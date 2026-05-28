"""Input loaders for manual text, TXT, CSV, PDF, and DOCX files."""

from __future__ import annotations

import io
import uuid
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from pypdf import PdfReader

from .preprocessing import detect_language_simple


STANDARD_COLUMNS = ["id", "source_name", "text", "language", "category"]
TEXT_COLUMN_GUESSES = ["text", "content", "body", "article", "document", "description"]


def _empty_documents() -> pd.DataFrame:
    return pd.DataFrame(columns=STANDARD_COLUMNS)


def _new_id(prefix: str = "doc") -> str:
    return f"{prefix}_{uuid.uuid4().hex[:10]}"


def _source_name(uploaded_file: object, fallback: str = "uploaded_file") -> str:
    return str(getattr(uploaded_file, "name", fallback) or fallback)


def _read_uploaded_bytes(uploaded_file: object) -> io.BytesIO:
    """Read a Streamlit UploadedFile into a fresh BytesIO buffer."""
    if hasattr(uploaded_file, "seek"):
        uploaded_file.seek(0)
    return io.BytesIO(uploaded_file.read())


def _document_frame(rows: Iterable[dict]) -> pd.DataFrame:
    frame = pd.DataFrame(list(rows))
    if frame.empty:
        return _empty_documents()
    for column in STANDARD_COLUMNS:
        if column not in frame.columns:
            frame[column] = None
    return frame[STANDARD_COLUMNS]


def _guess_text_column(columns: Iterable[str]) -> str | None:
    normalized = {str(column).strip().lower(): column for column in columns}
    for guess in TEXT_COLUMN_GUESSES:
        if guess in normalized:
            return str(normalized[guess])
    return None


def load_manual_text(text: str, category: str = "manual") -> pd.DataFrame:
    """Load pasted text into the standard document structure."""
    value = str(text or "").strip()
    if not value:
        return _empty_documents()

    return _document_frame(
        [
            {
                "id": _new_id("manual"),
                "source_name": "manual_input",
                "text": value,
                "language": detect_language_simple(value),
                "category": category or "manual",
            }
        ]
    )


def load_txt_file(uploaded_file: object, category: str = "unknown") -> pd.DataFrame:
    """Load a plain text upload into the standard document structure."""
    buffer = _read_uploaded_bytes(uploaded_file)
    raw = buffer.getvalue()

    try:
        text = raw.decode("utf-8")
    except UnicodeDecodeError:
        text = raw.decode("latin-1", errors="replace")

    source = _source_name(uploaded_file)
    if not text.strip():
        return _empty_documents()

    return _document_frame(
        [
            {
                "id": _new_id("txt"),
                "source_name": source,
                "text": text,
                "language": detect_language_simple(text),
                "category": category or "unknown",
            }
        ]
    )


def load_csv_file(
    uploaded_file: object,
    text_col: str | None = None,
    id_col: str | None = None,
    language_col: str | None = None,
    category_col: str | None = None,
    default_category: str = "unknown",
) -> pd.DataFrame:
    """Load a CSV upload and map it into the standard document structure."""
    buffer = _read_uploaded_bytes(uploaded_file)
    raw_df = pd.read_csv(buffer)

    if text_col is None:
        text_col = _guess_text_column(raw_df.columns)
    if id_col is None and "id" in raw_df.columns:
        id_col = "id"
    if language_col is None and "language" in raw_df.columns:
        language_col = "language"
    if category_col is None and "category" in raw_df.columns:
        category_col = "category"

    if not text_col or text_col not in raw_df.columns:
        guesses = ", ".join(TEXT_COLUMN_GUESSES)
        raise ValueError(
            f"Could not find a text column. Choose one explicitly or use one of: {guesses}."
        )

    source = _source_name(uploaded_file)
    rows = []

    for index, row in raw_df.iterrows():
        text = "" if pd.isna(row[text_col]) else str(row[text_col])
        if not text.strip():
            continue

        document_id = (
            str(row[id_col])
            if id_col and id_col in raw_df.columns and not pd.isna(row[id_col])
            else f"{Path(source).stem}_{index + 1}"
        )
        language = (
            str(row[language_col])
            if language_col and language_col in raw_df.columns and not pd.isna(row[language_col])
            else detect_language_simple(text)
        )
        category = (
            str(row[category_col])
            if category_col and category_col in raw_df.columns and not pd.isna(row[category_col])
            else default_category or "unknown"
        )

        rows.append(
            {
                "id": document_id,
                "source_name": source,
                "text": text,
                "language": language,
                "category": category,
            }
        )

    return _document_frame(rows)


def load_pdf_file(
    uploaded_file: object,
    category: str = "unknown",
    split_pages: bool = False,
) -> tuple[pd.DataFrame, str | None]:
    """Extract text from a PDF upload using pypdf."""
    source = _source_name(uploaded_file)
    buffer = _read_uploaded_bytes(uploaded_file)

    try:
        reader = PdfReader(buffer)
    except Exception as exc:
        return _empty_documents(), f"Could not read {source}: {exc}"

    page_rows = []
    collected_text = []

    for page_number, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            collected_text.append(page_text)
            if split_pages:
                page_rows.append(
                    {
                        "id": f"{Path(source).stem}_page_{page_number}",
                        "source_name": f"{source} page {page_number}",
                        "text": page_text,
                        "language": detect_language_simple(page_text),
                        "category": category or "unknown",
                    }
                )

    if not collected_text:
        return (
            _empty_documents(),
            f"No extractable text found in {source}. If this is a scanned PDF, OCR is not enabled.",
        )

    if split_pages:
        return _document_frame(page_rows), None

    text = "\n\n".join(collected_text)
    return (
        _document_frame(
            [
                {
                    "id": _new_id("pdf"),
                    "source_name": source,
                    "text": text,
                    "language": detect_language_simple(text),
                    "category": category or "unknown",
                }
            ]
        ),
        None,
    )


def load_docx_file(uploaded_file: object, category: str = "unknown") -> pd.DataFrame:
    """Extract paragraphs and table cells from a DOCX upload."""
    source = _source_name(uploaded_file)
    buffer = _read_uploaded_bytes(uploaded_file)
    document = Document(buffer)

    chunks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    chunks.append(cell.text)

    text = "\n".join(chunks).strip()
    if not text:
        return _empty_documents()

    return _document_frame(
        [
            {
                "id": _new_id("docx"),
                "source_name": source,
                "text": text,
                "language": detect_language_simple(text),
                "category": category or "unknown",
            }
        ]
    )


def load_multiple_files(
    uploaded_files: Iterable[object],
    category: str = "unknown",
) -> tuple[pd.DataFrame, list[str]]:
    """Load multiple uploads, routing each file by extension."""
    frames = []
    warnings: list[str] = []

    for uploaded_file in uploaded_files or []:
        source = _source_name(uploaded_file)
        extension = Path(source).suffix.lower()

        try:
            if extension == ".txt":
                frame = load_txt_file(uploaded_file, category=category)
            elif extension == ".csv":
                frame = load_csv_file(uploaded_file, default_category=category)
            elif extension == ".pdf":
                frame, warning = load_pdf_file(uploaded_file, category=category)
                if warning:
                    warnings.append(warning)
            elif extension == ".docx":
                frame = load_docx_file(uploaded_file, category=category)
            else:
                warnings.append(f"Unsupported file type for {source}.")
                continue

            if frame.empty:
                warnings.append(f"No usable text found in {source}.")
            else:
                frames.append(frame)
        except Exception as exc:
            warnings.append(f"Could not process {source}: {exc}")

    if not frames:
        return _empty_documents(), warnings

    return pd.concat(frames, ignore_index=True), warnings
